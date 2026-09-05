import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Page Margins (0.75 in)
for s in doc.sections:
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)

PRIMARY = RGBColor(5, 150, 105)     # #059669 Emerald
DARK = RGBColor(15, 23, 42)         # #0F172A Dark Slate
MUTED = RGBColor(100, 116, 139)     # #64748B Slate
ALERT_RED = RGBColor(185, 28, 28)   # #B91C1C Red

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_banner(title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(subtitle)
    run2.font.name = 'Arial'
    run2.font.size = Pt(11)
    run2.font.color.rgb = MUTED

def add_stage_header(stage_num, title, goal_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f'STAGE {stage_num}: {title}')
    run.font.name = 'Arial'
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = PRIMARY

    p_goal = doc.add_paragraph()
    p_goal.paragraph_format.space_before = Pt(0)
    p_goal.paragraph_format.space_after = Pt(5)
    run_goal = p_goal.add_run('Goal: ' + goal_text)
    run_goal.font.name = 'Arial'
    run_goal.font.size = Pt(9.5)
    run_goal.font.italic = True
    run_goal.font.color.rgb = MUTED

def add_checklist_item(step_label, title, details, success_signal=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    run_box = p.add_run('☐  ')
    run_box.font.name = 'Arial'
    run_box.font.size = Pt(11)
    run_box.font.bold = True
    run_box.font.color.rgb = PRIMARY
    
    run_title = p.add_run(f'{step_label}: {title}\n')
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK

    for d in details:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.left_indent = Inches(0.45)
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(1)
        r_bullet = p_sub.add_run('• ')
        r_bullet.font.bold = True
        r_bullet.font.color.rgb = MUTED
        r_text = p_sub.add_run(d)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(9)
        r_text.font.color.rgb = DARK

    if success_signal:
        p_succ = doc.add_paragraph()
        p_succ.paragraph_format.left_indent = Inches(0.45)
        p_succ.paragraph_format.space_before = Pt(1)
        p_succ.paragraph_format.space_after = Pt(4)
        r_succ_label = p_succ.add_run('✔ Expected Result: ')
        r_succ_label.font.name = 'Arial'
        r_succ_label.font.size = Pt(9)
        r_succ_label.font.bold = True
        r_succ_label.font.color.rgb = PRIMARY
        r_succ = p_succ.add_run(success_signal)
        r_succ.font.name = 'Arial'
        r_succ.font.size = Pt(9)
        r_succ.font.color.rgb = DARK

add_header_banner('Clinic Staff Testing Checklist', 'Animal Bite Treatment Center (ABTC / RHU) — Step-by-Step QA Manual')

p_meta = doc.add_paragraph()
p_meta.paragraph_format.space_after = Pt(10)
r = p_meta.add_run('Google Docs & Microsoft Word Edition • Designed for Doctors, Nurses, Clerks & QA Testers • URL: http://localhost:5173')
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = MUTED

# Table 1: Login Credentials
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

headers = ['Staff Desk', 'Login Email', 'Password', 'What This Role Tests']
widths = [Inches(1.5), Inches(2.2), Inches(1.3), Inches(2.0)]

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].width = widths[i]
    set_cell_background(hdr_cells[i], '059669')
    set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

cred_data = [
    ('Clinic Admin', 'admin@clinic.com', 'password123', 'Sets clinic open days & checks vaccine stock.'),
    ('Registration Clerk', 'registration@clinic.com', 'password123', 'Enrolls walk-in bite victims & issues queue ticket.'),
    ('Doctor (Triage)', 'triage@clinic.com', 'password123', 'Examines bite, grades rabies risk & auto-refers.'),
    ('Nurse (Treatment)', 'treatment@clinic.com', 'password123', 'Injects Day 0, records batch & checks return dates.')
]

for row_idx, row in enumerate(cred_data):
    row_cells = table.add_row().cells
    bg = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
    for i, val in enumerate(row):
        row_cells[i].text = val
        row_cells[i].width = widths[i]
        set_cell_background(row_cells[i], bg)
        set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)
        for p in row_cells[i].paragraphs:
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                r.font.color.rgb = DARK

# STAGE 1: Admin
add_stage_header(1, 'Clinic Administrator (Schedule & Vaccine Setup)', 'Verify clinic opening days and ensure rabies vaccine vials are in stock.')
add_checklist_item('Step 1.1', 'Sign In as Clinic Admin', 
    ['Go to http://localhost:5173/login', 'Type admin@clinic.com and password123. Click Sign In.'],
    'Admin dashboard loads successfully.')
add_checklist_item('Step 1.2', 'Verify Clinic Operating Days', 
    ['In the left sidebar, click Clinic Setup -> Operating Schedule (or /setup/schedule).',
     'Verify Tuesday to Saturday are marked OPEN.',
     'Verify Sunday and Monday are marked CLOSED.'],
    'The clinic operates Tue-Sat; system will NEVER book follow-ups on Sun/Mon.')
add_checklist_item('Step 1.3', 'Check Vaccine Inventory Stock', 
    ['In the left sidebar, click Inventory (/inventory).',
     'Verify at least one brand (e.g. Speeda or Verorab) shows available stock.'],
    'Stock badge displays Green or Active vials.')

# STAGE 2: Registration
add_stage_header(2, 'Registration Desk (Patient Arrival & Form 1)', 'Encode a walk-in bite victim and issue their automated queue ticket.')
add_checklist_item('Step 2.1', 'Sign In as Registration Staff', 
    ['Go to http://localhost:5173/login', 'Type registration@clinic.com and password123. Click Sign In.'],
    'Registration desk loads.')
add_checklist_item('Step 2.2', 'Add New Bite Patient', 
    ['Click Patients in the left menu, then click the green "+ Add Patient" button.',
     'First Name: Juan | Last Name: Dela Cruz | Sex: Male',
     'Date of Birth: 1995-05-15 (Verify Age auto-calculates to 31)',
     'Mobile Number: 09171234567 (11 digits) | Emergency Contact: Maria Dela Cruz / 09181234567',
     'Address: Tagoloan, Poblacion, Zone 2',
     'Queue Category: Regular (or Senior/Pregnant/PWD) | Priority: Normal',
     'Click "Save Patient Record".'],
    'Green notification: "Patient enrolled successfully". Juan Dela Cruz appears in Patients list.')
add_checklist_item('Step 2.3', 'Verify Automatic Queue Ticket', 
    ['Click Queue Dashboard in the sidebar (/queue).',
     'Verify Juan Dela Cruz appears on the board with Ticket # (e.g. #1).',
     'Type: New Case | Status: Waiting.'],
    'Ticket is actively counting waiting time on the triage queue screen.')

# STAGE 3: Doctor Triage
add_stage_header(3, 'Doctor Triage Desk (Examine Wound & Prescribe Form 2)', 'Examine wound, assign exposure Category (I-III), prescribe PEP, and trigger automated handoff.')
add_checklist_item('Step 3.1', 'Sign In as Doctor & Open Patient', 
    ['Log in as triage@clinic.com / password123.',
     'Go to Queue Dashboard (/queue). Find Juan Dela Cruz.',
     'Click the green "View" button on his row.'],
    'Pop-up window opens right over the queue board WITHOUT leaving or reloading the page!')
add_checklist_item('Step 3.2', 'Verify Protection Between Desks', 
    ['Tab 1 (Demographics): Shows Juan\'s personal details in read-only format.',
     'Tab 3 (Nurse Treatment): Displays amber lock: "Awaiting Doctor Triage (Form 2 Required)".'],
    'Nurse is blocked from injecting vaccine until physician completes clinical assessment.')
add_checklist_item('Step 3.3', 'Fill Form 2 & Save Consultation', 
    ['Click Tab 2 (Form 2 Doctor Triage).',
     'Weight: 65 kg | History: Dog bite on right forearm 2 hours ago.',
     'Animal: Dog | Animal Status: Alive / Under Observation | Bite Type: Bite',
     'Exposure Category: Category III (Severe bite) | Washed with soap: Yes',
     'Tetanus Toxoid: Given | Vaccine Regimen: Intradermal (ID) - Option A',
     'ERIG (Anti-Rabies Serum): Prescribed (Calculates 65 kg x 40 IU = 2600 IU).',
     'Click "Save Consultation" / "Save Record".'],
    'Green toast: "Consultation record saved successfully. Patient referred to Treatment."')
add_checklist_item('Step 3.4', 'Verify Automated Handoff (The Magic!)', 
    ['Look at Juan\'s row on the Queue Dashboard.',
     'Ticket automatically changed Type: New Case -> Vaccination!',
     'Status reset to Waiting for the nurse.',
     'Notes state: "Doctor completed Form 2 — referred to Treatment."'],
    'Doctor completed referral with ZERO manual queue transfer clicks!')

# STAGE 4: Nurse Treatment
add_stage_header(4, 'Nurse Treatment Desk (Day 0 Shot & Option A Calendar)', 'Inject Day 0, deduct stock, complete queue ticket, and verify auto-scheduled return dates.')
add_checklist_item('Step 4.1', 'Sign In as Nurse & Open Patient Record', 
    ['Log in as treatment@clinic.com / password123.',
     'Go to Queue Dashboard (/queue). Locate Juan Dela Cruz under Treatment Queue.',
     'Click the green "View" button.'],
    'In-place window opens. Tab 2 (Doctor Triage) is locked to prevent tampering.')
add_checklist_item('Step 4.2', 'Administer Day 0 Shot', 
    ['Click Tab 3 (Nurse Treatment). Lock is gone!',
     'Dose 0 (Day 0) row has emerald badge: "Ready to Administer".',
     'Select Vaccine Brand: Speeda (or active brand) | Pick Batch Number.',
     'Route: ID (Intradermal) | Site: Left Deltoid (0.1mL) & Right Deltoid (0.1mL).',
     'ERIG section: Check Administered (if prescribed).',
     'Click "Save Vaccination Record".'],
    'Dose 0 immediately turns into "Administered" and becomes PERMANENTLY LOCKED.')
add_checklist_item('Step 4.3', 'Verify Ticket Completion & Stock Deduction', 
    ['Check Queue Dashboard: Juan\'s ticket is marked "Completed" and archived.',
     'Check Inventory (/inventory): 1 vial of vaccine is subtracted from stock.'],
    'Ticket clears off active board and inventory count decrements accurately.')
add_checklist_item('Step 4.4', 'Verify Weekend Shift Rule (Option A Resolution)', 
    ['Go to Vaccination Schedule in sidebar (/vaccinations/schedule). Find Juan Dela Cruz.',
     'Day 0: Completed (Green dot) | Day 3: Scheduled (+3 days from Day 0).',
     'CRITICAL CHECK: If Day 3 lands on a closed Sun/Mon, verify amber tag: "Shifted +1d/+2d — clinic closed on weekend" and shifts date to Tuesday!',
     'Day 7 (+7d) and Day 28 (+28d) are auto-scheduled.'],
    'Zero appointments booked on closed clinic days!')

# STAGE 5: Follow-Up
add_stage_header(5, 'Routine Follow-Up Visit (Day 3 Return)', 'Returning follow-up patient checks in directly at Nurse Desk, and past shots cannot be modified.')
add_checklist_item('Step 5.1', 'Direct Nurse Station Check-In', 
    ['Patient arrives on Day 3.',
     'Nurse opens Nurse Desk (/nurse/patients) -> clicks "Due Today" tab.',
     'Finds Juan Dela Cruz -> clicks "Check In" button and confirms.'],
    'Patient gets ticket directly for Treatment Desk. Completely bypasses Registration and Doctor!')
add_checklist_item('Step 5.2', 'Record Day 3 Shot (Hybrid Immutability)', 
    ['Click "Record Dose" (or open on /queue).',
     'Day 0 row is DISABLED/LOCKED (brand, batch, date cannot be altered).',
     'Day 3 row is ACTIVE with "Ready to Administer" badge.',
     'Future doses (Day 7, 28) show "Pending" without false red stock errors.',
     'Brand automatically pre-selects Speeda to match Day 0.',
     'Select Day 3 Batch Number, enter Site, and click "Save Vaccination Record".'],
    'Day 3 is saved and locked. Ticket completes automatically. Day 7 is next active dose.')

# STAGE 6: Special Cases
add_stage_header(6, 'Special Clinical Cases to Test', 'Special patient situations: Boosters, External clinic transfers, and Second Chance queue.')
add_checklist_item('Case 6.1', 'Priority Patients (Seniors, Pregnant, PWD)', 
    ['Register patient with Category: Senior Citizen, Pregnant, or PWD.',
     'Check on Queue Dashboard: Shows colored badge and ranks ahead of regular tickets.'],
    'Priority queuing works accurately.')
add_checklist_item('Case 6.2', 'Patient Does Not Hear Name (Second Chance Queue)', 
    ['On /queue, click the purple UserBlock icon ("No Response").',
     'Ticket moves into Second Chance Queue with 10-minute timer.',
     'Click "Recall" to call them again, or "Mark Absent" if they left.'],
    'Queue second chance workflow functions cleanly.')
add_checklist_item('Case 6.3', 'Booster Protocol (Patient Bit Years Later)', 
    ['Doctor checks "Prior Immunization Verified" and selects "2-Dose Booster".',
     'Verify ERIG checkbox is DISABLED with notice: "Contraindicated: Prior immunization verified".',
     'Verify Form 3 creates Day 0 and Day 3 ONLY (Days 7 and 28 are omitted).',
     'Administer Day 3 -> Episode marks "Regimen Completed".'],
    'Booster regimen strictly enforces 2 doses with no ERIG.')
add_checklist_item('Case 6.4', 'Dose Given at Another Hospital (Transferred-In)', 
    ['In Form 3 on Day 0, check "Transferred-In (External Clinic)".',
     'Enter hospital name (e.g. Provincial Hospital).',
     'Status displays neutral "External Clinic" badge (NOT red Missing Stock Info error).',
     'Saving form does NOT fail or deduct local stock for Day 0.'],
    'External clinic doses record smoothly without local inventory errors.')

# STAGE 7: Safety Rules
add_stage_header(7, 'Clinical Safety Red Flags (What Must Be Blocked)', 'Ensure system actively prevents dangerous medical-legal errors.')

table2 = doc.add_table(rows=1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.autofit = False

headers2 = ['Clinical Hazard to Block', 'How to Test', 'Expected Protection']
widths2 = [Inches(2.2), Inches(2.8), Inches(2.0)]

hdr_cells2 = table2.rows[0].cells
for i, h in enumerate(headers2):
    hdr_cells2[i].text = h
    hdr_cells2[i].width = widths2[i]
    set_cell_background(hdr_cells2[i], 'B91C1C')
    set_cell_margins(hdr_cells2[i], top=120, bottom=120, left=150, right=150)
    for p in hdr_cells2[i].paragraphs:
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

hazard_data = [
    ('Tampering with Doctor Diagnosis', 'Open Form 2 after Day 0 is injected.', 'All fields locked/read-only. Doctors can only add progress addendums.'),
    ('Overwriting Administered Shots', 'Open Form 3 for patient with Day 0 on file.', 'Day 0 brand, batch, date, and nurse initials are completely disabled.'),
    ('Booking on Closed Weekend Days', 'Check appointments on /vaccinations/schedule.', 'Zero appointments on Sunday or Monday; shifted forward to Tuesday.'),
    ('Giving ERIG to Booster Patient', 'Select 2-Dose Booster regimen in Form 2.', 'ERIG checkbox is disabled and flagged contraindicated.'),
    ('Non-Admin Altering Legal Name', 'Open demographic edit modal as Nurse or Clerk.', 'Legal Name and Birthdate are locked with padlock icons (Admin only).'),
    ('Losing Queue Screen on View', 'Click "View" on any row in Queue Dashboard.', 'Opens in-place dialog window; stays on /queue URL.')
]

for row_idx, row in enumerate(hazard_data):
    row_cells = table2.add_row().cells
    bg = 'FEF2F2' if row_idx % 2 == 1 else 'FFFFFF'
    for i, val in enumerate(row):
        row_cells[i].text = val
        row_cells[i].width = widths2[i]
        set_cell_background(row_cells[i], bg)
        set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)
        for p in row_cells[i].paragraphs:
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                r.font.color.rgb = DARK

# Tester Sign-Off Sheet
p_sign = doc.add_paragraph()
p_sign.paragraph_format.space_before = Pt(18)
p_sign.paragraph_format.space_after = Pt(4)
r_s = p_sign.add_run('Tester Sign-Off Sheet')
r_s.font.name = 'Arial'
r_s.font.size = Pt(12)
r_s.font.bold = True
r_s.font.color.rgb = PRIMARY

sign_lines = [
    'Tester Name: _________________________________________',
    'Date Tested: _________________________________________',
    'Browser Used: [  ] Google Chrome   [  ] Microsoft Edge   [  ] Firefox',
    'Overall Result: [  ] ALL TESTS PASSED    [  ] ISSUES FOUND',
    '',
    '☐ Registration created patient and issued sequential queue ticket.',
    '☐ Doctor triage saved Form 2 and auto-referred ticket to Treatment Desk.',
    '☐ Nurse administered Day 0 and auto-completed queue ticket.',
    '☐ Vaccine stock decremented accurately in Inventory.',
    '☐ PEP Option A appointments generated without landing on closed days.',
    '☐ Returning follow-up patient checked in directly at Nurse Desk.',
    '☐ Past administered doses remained completely immutable.',
    '☐ Queue patient details opened smoothly in the in-place modal dialog.'
]

for line in sign_lines:
    p_l = doc.add_paragraph()
    p_l.paragraph_format.space_before = Pt(2)
    p_l.paragraph_format.space_after = Pt(2)
    r_l = p_l.add_run(line)
    r_l.font.name = 'Arial'
    r_l.font.size = Pt(9.5)
    r_l.font.color.rgb = DARK

output_path = os.path.abspath('docs/CLINIC_STAFF_TESTING_CHECKLIST.docx')
doc.save(output_path)
print(f'SUCCESS: DOCX created at {output_path}')